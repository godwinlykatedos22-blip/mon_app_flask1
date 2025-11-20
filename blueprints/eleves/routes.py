# blueprints/eleves/routes.py

from flask import Blueprint, render_template, redirect, url_for, flash, request, send_file, jsonify
from flask_login import login_required
from models import db, Student, Classe, Parent
from forms import StudentForm, DeleteForm
import os
from werkzeug.utils import secure_filename
from openpyxl import load_workbook, Workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

# Blueprint principal pour la gestion des élèves
eleves_bp = Blueprint(
    "eleves",
    __name__,
    template_folder="../../templates"
)


# --------------------------------------------------------
# LISTE DES ÉLÈVES
# --------------------------------------------------------

@eleves_bp.route("/")
@login_required
def list_eleves():
    students = Student.query.order_by(Student.last_name).all()
    delete_form = DeleteForm()  # Crée le formulaire pour le CSRF
    return render_template("eleves/eleves_list.html", students=students, delete_form=delete_form)


# --------------------------------------------------------
# OBTENIR L'ID D'UNE CLASSE PAR SON NOM
# --------------------------------------------------------
@eleves_bp.route("/get-class-id/<class_name>", methods=["GET"])
@login_required
def get_class_id(class_name):
    """Obtient l'ID d'une classe par son nom, utilisé pour l'export dynamique"""
    classe = Classe.query.filter_by(name=class_name).first()
    if classe:
        return jsonify({"class_id": classe.id})
    else:
        # Chercher les élèves avec cette classe (créer la classe si besoin)
        student = Student.query.join(Classe).filter(Classe.name == class_name).first()
        if student and student.classe:
            return jsonify({"class_id": student.classe.id})
        return jsonify({"class_id": None}), 404


# --------------------------------------------------------
# AJOUT D'UN ÉLÈVE
# --------------------------------------------------------
@eleves_bp.route("/add", methods=["GET", "POST"])
@login_required
def add_eleve():
    form = StudentForm()
    form.class_id.choices = [(c.id, c.name) for c in Classe.query.all()]

    if form.validate_on_submit():

        # 1. Création de l'élève
        student = Student(
            first_name=form.first_name.data,
            last_name=form.last_name.data,
            birthdate=form.birthdate.data,
            class_id=form.class_id.data,
        )
        db.session.add(student)
        db.session.commit()  # On commit pour obtenir student.id

        # 2. Création ou récupération du parent SI champs renseignés
        if form.parent_first_name.data or form.parent_last_name.data or form.parent_phone.data:
            # Chercher un parent existant avec le même numéro de téléphone
            parent = None
            if form.parent_phone.data:
                parent = Parent.query.filter_by(phone_e164=form.parent_phone.data).first()
            
            # Si pas trouvé, chercher par nom+prénom
            if not parent and form.parent_first_name.data and form.parent_last_name.data:
                parent = Parent.query.filter_by(
                    first_name=form.parent_first_name.data.strip(),
                    last_name=form.parent_last_name.data.strip()
                ).first()
            
            # Si toujours pas trouvé, créer un nouveau parent
            if not parent:
                parent = Parent(
                    first_name=form.parent_first_name.data,
                    last_name=form.parent_last_name.data,
                    phone_e164=form.parent_phone.data,
                    whatsapp_optin=form.parent_whatsapp.data
                )
                db.session.add(parent)
                db.session.commit()
            else:
                # Mettre à jour les infos du parent existant si nécessaire
                if form.parent_first_name.data:
                    parent.first_name = form.parent_first_name.data
                if form.parent_last_name.data:
                    parent.last_name = form.parent_last_name.data
                if form.parent_phone.data:
                    parent.phone_e164 = form.parent_phone.data
                parent.whatsapp_optin = form.parent_whatsapp.data
                db.session.commit()

            # 3. Association parent ↔ élève (si pas déjà lié)
            if student not in parent.students:
                parent.students.append(student)
                db.session.commit()

        flash("Élève et parent(s) enregistrés avec succès ! 🎉", "success")
        return redirect(url_for("eleves.list_eleves"))

    # Si le formulaire a été posté mais n'est pas valide, afficher les erreurs
    if request.method == "POST" and form.errors:
        for field, errors in form.errors.items():
            for err in errors:
                flash(f"Erreur champ {field} : {err}", "danger")

    return render_template("eleves/eleves_form.html", form=form)



# --------------------------------------------------------
# MODIFICATION D'UN ÉLÈVE
# --------------------------------------------------------
@eleves_bp.route("/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_eleve(id):
    student = Student.query.get_or_404(id)
    form = StudentForm(obj=student)

    form.class_id.choices = [(c.id, c.name) for c in Classe.query.all()]

    if form.validate_on_submit():

        # Mise à jour des infos étudiant
        student.first_name = form.first_name.data
        student.last_name = form.last_name.data
        student.birthdate = form.birthdate.data
        student.class_id = form.class_id.data

        # Gestion du parent
        parent = student.parents[0] if student.parents else None

        if form.parent_first_name.data or form.parent_last_name.data or form.parent_phone.data:

            if not parent:
                parent = Parent()
                student.parents.append(parent)

            parent.first_name = form.parent_first_name.data
            parent.last_name = form.parent_last_name.data
            parent.phone_e164 = form.parent_phone.data
            parent.whatsapp_optin = form.parent_whatsapp.data
        
        form.populate_obj(student)
        db.session.commit()
        flash("Élève mis à jour avec succès 🎉", "success")
        return redirect(url_for("eleves.list_eleves"))

    # Si POST invalide, afficher erreurs
    if request.method == "POST" and form.errors:
        for field, errors in form.errors.items():
            for err in errors:
                flash(f"Erreur champ {field} : {err}", "danger")

    return render_template("eleves/eleves_form.html", form=form)



@eleves_bp.route("/delete/<int:id>", methods=["POST"])
@login_required
def delete_eleve(id):
    student = Student.query.get_or_404(id)
    db.session.delete(student)
    db.session.commit()
    flash("Élève supprimé 🗑️", "info")
    return redirect(url_for("eleves.list_eleves"))


@eleves_bp.route("/import", methods=["GET", "POST"])
@login_required
def import_eleves():
    classes = Classe.query.order_by(Classe.name).all()

    if request.method == "POST":
        file = request.files.get("excel_file")

        if not file:
            flash("Veuillez sélectionner un fichier Excel.", "danger")
            return redirect(url_for("eleves.import_eleves"))

        filename = secure_filename(file.filename)

        if not filename.endswith(".xlsx"):
            flash("Le fichier doit être au format .xlsx", "danger")
            return redirect(url_for("eleves.import_eleves"))

        os.makedirs("uploads", exist_ok=True)
        filepath = os.path.join("uploads", filename)
        file.save(filepath)

        try:
            wb = load_workbook(filepath)
            sheet = wb.active

            count = 0

            for row in sheet.iter_rows(min_row=2, values_only=True):
                # Récupérer les colonnes : nom, prénom, date, classe, parent_prenom, parent_nom, parent_phone, whatsapp
                if len(row) < 4:
                    continue
                
                last_name = row[0]
                first_name = row[1]
                birthdate = row[2] if len(row) > 2 else None
                classe_name = row[3] if len(row) > 3 else None
                parent_first_name = row[4] if len(row) > 4 else None
                parent_last_name = row[5] if len(row) > 5 else None
                parent_phone = row[6] if len(row) > 6 else None
                parent_whatsapp = row[7] if len(row) > 7 else None

                if not last_name or not first_name:
                    continue

                # Vérifier / créer classe
                classe = Classe.query.filter_by(name=str(classe_name).strip()).first()
                if not classe:
                    classe = Classe(name=str(classe_name).strip())
                    db.session.add(classe)
                    db.session.commit()

                # Vérifier doublon
                exists = Student.query.filter_by(
                    last_name=str(last_name).strip(),
                    first_name=str(first_name).strip(),
                    class_id=classe.id
                ).first()

                if exists:
                    continue

                student = Student(
                    last_name=str(last_name).strip(),
                    first_name=str(first_name).strip(),
                    birthdate=birthdate if isinstance(birthdate, (str, type(None))) else birthdate.date() if hasattr(birthdate, 'date') else None,
                    class_id=classe.id
                )

                db.session.add(student)
                db.session.flush()  # Flush pour obtenir l'ID sans commit

                # Gérer le parent
                if parent_first_name or parent_last_name or parent_phone:
                    parent = None
                    
                    # Chercher parent par téléphone
                    if parent_phone:
                        parent_phone_str = str(parent_phone).strip()
                        parent = Parent.query.filter_by(phone_e164=parent_phone_str).first()
                    
                    # Chercher par nom+prénom
                    if not parent and parent_first_name and parent_last_name:
                        parent = Parent.query.filter_by(
                            first_name=str(parent_first_name).strip(),
                            last_name=str(parent_last_name).strip()
                        ).first()
                    
                    # Créer nouveau parent
                    if not parent:
                        parent = Parent(
                            first_name=str(parent_first_name).strip() if parent_first_name else None,
                            last_name=str(parent_last_name).strip() if parent_last_name else None,
                            phone_e164=str(parent_phone).strip() if parent_phone else None,
                            whatsapp_optin=(str(parent_whatsapp).lower() in ['oui', 'yes', 'true', '1']) if parent_whatsapp else False
                        )
                        db.session.add(parent)
                        db.session.flush()
                    
                    # Lier parent à élève
                    if student not in parent.students:
                        parent.students.append(student)

                count += 1

            db.session.commit()
            flash(f"{count} élèves importés avec succès !", "success")

        except Exception as e:
            db.session.rollback()
            flash(f"Erreur lors de l'import : {str(e)}", "danger")

        return redirect(url_for("eleves.list_eleves"))

    return render_template("eleves/eleves_import.html", classes=classes)


@eleves_bp.route("/download-template")
@login_required
def download_template():
    classes = Classe.query.order_by(Classe.name).all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Élèves"

    # En-têtes avec colonnes parent
    ws["A1"] = "Nom"
    ws["B1"] = "Prénom"
    ws["C1"] = "Date de naissance"
    ws["D1"] = "Classe"
    ws["E1"] = "Prénom Parent/Tuteur"
    ws["F1"] = "Nom Parent/Tuteur"
    ws["G1"] = "Téléphone Parent (E.164)"
    ws["H1"] = "WhatsApp (Oui/Non)"

    # Exemple de données
    ws["A2"] = "Dupont"
    ws["B2"] = "Jean"
    ws["C2"] = "12/05/2008"
    ws["D2"] = "6ème"
    ws["E2"] = "Marie"
    ws["F2"] = "Dupont"
    ws["G2"] = "+2290196969696"
    ws["H2"] = "Oui"

    # Style en-têtes
    from openpyxl.styles import Font, PatternFill
    header_fill = PatternFill(start_color="003F7F", end_color="003F7F", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    
    for cell in ws[1]:
        if cell.value:
            cell.fill = header_fill
            cell.font = header_font

    # Feuille classes disponibles
    ws2 = wb.create_sheet("Classes disponibles")
    ws2["A1"] = "Nom de la classe"
    ws2["A1"].fill = header_fill
    ws2["A1"].font = header_font

    row = 2
    for c in classes:
        ws2[f"A{row}"] = c.name
        row += 1

    # Feuille instructions
    ws3 = wb.create_sheet("Instructions")
    ws3["A1"] = "INSTRUCTIONS D'IMPORT"
    ws3["A1"].font = Font(bold=True, size=12)
    
    instructions = [
        "",
        "Remplissez la feuille 'Élèves' avec les colonnes suivantes :",
        "",
        "• Nom : Nom de famille de l'élève (OBLIGATOIRE)",
        "• Prénom : Prénom de l'élève (OBLIGATOIRE)",
        "• Date de naissance : Format YYYY-MM-DD (optionnel)",
        "• Classe : Nom de la classe (voir feuille 'Classes disponibles')",
        "• Prénom Parent : Prénom du parent/tuteur (optionnel)",
        "• Nom Parent : Nom du parent/tuteur (optionnel)",
        "• Téléphone Parent : Format E.164 ex: +33612345678 (optionnel)",
        "• WhatsApp : Oui ou Non (optionnel)",
        "",
        "Notes importantes :",
        "- Si un parent avec le même téléphone existe déjà, l'élève sera lié au parent existant",
        "- Les doublons d'élèves (même nom, prénom, classe) ne seront pas importés",
        "- Assurez-vous que les noms de classe correspondent exactement",
    ]

    for idx, instruction in enumerate(instructions, start=2):
        ws3[f"A{idx}"] = instruction

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name="modele_import_eleves.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# --------------------------------------------------------
# TÉLÉCHARGER EXEMPLE IMPORT AVEC DONNÉES
# --------------------------------------------------------
@eleves_bp.route("/download-example")
@login_required
def download_example():
    """Télécharge un fichier Excel d'exemple avec données"""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    
    classes = Classe.query.order_by(Classe.name).all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Élèves"

    # Largeurs de colonnes
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 12
    ws.column_dimensions['E'].width = 18
    ws.column_dimensions['F'].width = 18
    ws.column_dimensions['G'].width = 22
    ws.column_dimensions['H'].width = 15

    # Styles
    header_fill = PatternFill(start_color="003F7F", end_color="003F7F", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # En-têtes
    headers = ["Nom", "Prénom", "Date de naissance", "Classe", "Prénom Parent/Tuteur", "Nom Parent/Tuteur", "Téléphone Parent (E.164)", "WhatsApp (Oui/Non)"]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border

    # Données d'exemple
    example_data = [
        ["Dupont", "Jean", "2008-05-15", "6ème A", "Marie", "Dupont", "+33612345678", "Oui"],
        ["Dupont", "Sophie", "2010-03-10", "6ème A", "Marie", "Dupont", "+33612345678", "Oui"],
        ["Martin", "Alice", "2009-01-20", "6ème A", "Pierre", "Martin", "+33687654321", "Non"],
        ["Bernard", "Luc", "2008-11-05", "6ème A", "Anne", "Bernard", "+33698765432", "Oui"],
        ["Lefevre", "Emma", "2009-07-22", "6ème A", "Marc", "Lefevre", "+33645123789", "Oui"],
        ["Lefevre", "Thomas", "2010-09-12", "6ème A", "Marc", "Lefevre", "+33645123789", "Oui"],
    ]

    # Ajouter les données
    for row_num, row_data in enumerate(example_data, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num)
            cell.value = value
            cell.border = border
            
            if col_num in [4, 8]:  # Classe et WhatsApp
                cell.alignment = Alignment(horizontal='center')

    # Feuille Classes disponibles
    ws2 = wb.create_sheet("Classes disponibles")
    ws2.column_dimensions['A'].width = 20

    cell = ws2['A1']
    cell.value = "Nom de la classe"
    cell.fill = header_fill
    cell.font = header_font
    cell.border = border

    for idx, classe in enumerate(classes, 2):
        cell = ws2.cell(row=idx, column=1)
        cell.value = classe.name
        cell.border = border

    # Feuille Instructions (résumée)
    ws3 = wb.create_sheet("Instructions")
    ws3.column_dimensions['A'].width = 100

    title = ws3['A1']
    title.value = "INSTRUCTIONS D'IMPORT"
    title.font = Font(bold=True, size=12, color="FFFFFF")
    title.fill = header_fill

    instructions = [
        "",
        "✅ Colonnes obligatoires : Nom, Prénom, Classe",
        "",
        "📝 Colonnes optionnelles : Date naissance, Parent Prénom, Parent Nom, Téléphone, WhatsApp",
        "",
        "⚠️  Les parents sont regroupés par NUMÉRO DE TÉLÉPHONE",
        "    Si 2 élèves ont le même numéro, ils seront liés au même parent",
        "",
        "📞 Format téléphone : +33612345678 (E.164)",
        "    Exemple France : remplacez le 0 par 33",
        "",
        "📅 Format date : YYYY-MM-DD (ex: 2008-05-15)",
        "",
        "💬 WhatsApp : accepte Oui/Non/Yes/No/True/False/1/0",
    ]

    for idx, instruction in enumerate(instructions, 2):
        cell = ws3.cell(row=idx, column=1)
        cell.value = instruction
        cell.alignment = Alignment(horizontal='left', wrap_text=True)

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name="exemple_import_eleves.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# --------------------------------------------------------
# VOIR LES DÉTAILS D'UN ÉLÈVE
# --------------------------------------------------------
@eleves_bp.route("/details/<int:id>")
@login_required
def details_eleve(id):
    student = Student.query.get_or_404(id)
    delete_form = DeleteForm()
    return render_template("eleves/eleves_details.html", student=student, delete_form=delete_form)



# --------------------------------------------------------
# EXPORTER LA LISTE DES ÉLÈVES PAR CLASSE (EXCEL)
# --------------------------------------------------------
@eleves_bp.route("/export/excel/<int:class_id>")
@login_required
def export_excel(class_id):
    classe = Classe.query.get_or_404(class_id)
    students = Student.query.filter_by(class_id=class_id).order_by(Student.last_name).all()

    wb = Workbook()
    ws = wb.active
    ws.title = classe.name

    # En-têtes
    ws["A1"] = "ID"
    ws["B1"] = "Nom"
    ws["C1"] = "Prénom"
    ws["D1"] = "Date de naissance"
    ws["E1"] = "Parents"

    # Données
    row = 2
    for student in students:
        parents_names = ", ".join([f"{p.first_name} {p.last_name}" for p in student.parents])
        ws[f"A{row}"] = student.id
        ws[f"B{row}"] = student.last_name
        ws[f"C{row}"] = student.first_name
        ws[f"D{row}"] = str(student.birthdate) if student.birthdate else ""
        ws[f"E{row}"] = parents_names
        row += 1

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name=f"eleves_{classe.name}.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# --------------------------------------------------------
# EXPORTER LA LISTE DES ÉLÈVES PAR CLASSE (PDF)
# --------------------------------------------------------
@eleves_bp.route("/export/pdf/<int:class_id>")
@login_required
def export_pdf(class_id):
    classe = Classe.query.get_or_404(class_id)
    students = Student.query.filter_by(class_id=class_id).order_by(Student.last_name).all()

    stream = io.BytesIO()
    doc = SimpleDocTemplate(stream, pagesize=A4)
    elements = []

    # Titre
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#003f7f'),
        spaceAfter=12,
        alignment=1  # Center
    )
    elements.append(Paragraph(f"Liste des Élèves - Classe {classe.name}", title_style))
    elements.append(Spacer(1, 0.3 * inch))

    # Tableau
    data = [["ID", "Nom", "Prénom", "Date de naissance", "Parents"]]
    for student in students:
        parents_names = ", ".join([f"{p.first_name} {p.last_name}" for p in student.parents])
        data.append([
            str(student.id),
            student.last_name,
            student.first_name,
            str(student.birthdate) if student.birthdate else "",
            parents_names
        ])

    table = Table(data, colWidths=[0.5*inch, 1.2*inch, 1.2*inch, 1.5*inch, 2*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003f7f')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 0.3 * inch))

    # Pied de page
    footer_text = f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=1
    )
    elements.append(Paragraph(footer_text, footer_style))

    doc.build(elements)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name=f"eleves_{classe.name}.pdf",
        mimetype="application/pdf"
    )


# --------------------------------------------------------
# EXPORTER LA LISTE DES ÉLÈVES PAR CLASSE (WORD)
# --------------------------------------------------------
@eleves_bp.route("/export/word/<int:class_id>")
@login_required
def export_word(class_id):
    classe = Classe.query.get_or_404(class_id)
    students = Student.query.filter_by(class_id=class_id).order_by(Student.last_name).all()

    doc = Document()

    # Titre
    title = doc.add_heading(f"Liste des Élèves - Classe {classe.name}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.font.size = Pt(16)
    title_format.font.color.rgb = RGBColor(0, 63, 127)

    doc.add_paragraph()

    # Tableau
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'

    # En-têtes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Nom"
    hdr_cells[2].text = "Prénom"
    hdr_cells[3].text = "Date de naissance"
    hdr_cells[4].text = "Parents"

    # Données
    for student in students:
        row_cells = table.add_row().cells
        parents_names = ", ".join([f"{p.first_name} {p.last_name}" for p in student.parents])
        row_cells[0].text = str(student.id)
        row_cells[1].text = student.last_name
        row_cells[2].text = student.first_name
        row_cells[3].text = str(student.birthdate) if student.birthdate else ""
        row_cells[4].text = parents_names

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0]
    footer_format.font.size = Pt(8)
    footer_format.font.color.rgb = RGBColor(128, 128, 128)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    return send_file(
        stream,
        as_attachment=True,
        download_name=f"eleves_{classe.name}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

